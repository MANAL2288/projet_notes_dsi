from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io, os, re, json

# Chemin vers les logos (même dossier que utils.py)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO1 = os.path.join(BASE_DIR, "logo1.jpg")
LOGO2 = os.path.join(BASE_DIR, "logo2.jpg")

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), '000000')
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def generate_docx(data: dict) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)

    # HEADER — logo ministère
    header = doc.sections[0].header
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)
    htable = header.add_table(1, 1, width=Inches(6.3))
    htable.alignment = WD_TABLE_ALIGNMENT.CENTER
    hpara = htable.rows[0].cells[0].paragraphs[0]
    hpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hpara.add_run().add_picture(LOGO1, width=Cm(14))

    # FOOTER — logo génération green + adresse
    footer = doc.sections[0].footer
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)
    ftable = footer.add_table(1, 1, width=Inches(6.3))
    ftable.alignment = WD_TABLE_ALIGNMENT.CENTER
    fpara = ftable.rows[0].cells[0].paragraphs[0]
    fpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fpara.add_run().add_picture(LOGO2, width=Cm(16))

    # DSI + DATE
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(0)
    r_dsi = p1.add_run("DSI")
    r_dsi.bold = True
    r_dsi.font.size = Pt(11)
    r_dsi.add_tab()
    r_date = p1.add_run(f"Rabat, le {datetime.now().strftime('%d/%m/%Y')}")
    r_date.font.size = Pt(11)
    pPr = p1._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '8640')
    tabs.append(tab)
    pPr.append(tabs)

    doc.add_paragraph()

    # TITRE
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("Note")
    r_title.bold = True
    r_title.font.size = Pt(13)
    r_title.font.underline = True

    # TABLEAU
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows_data = [
        ("Cadre",                  data.get("cadre", ""),           False),
        ("Objet",                  data.get("objet", ""),           False),
        ("Participants ",          data.get("participants", ""),     False),
        ("Descriptif",             data.get("descriptif", ""),       True),
        ("Prochaine Action",       data.get("prochaine_action", ""), True),
        ("Instruction Sollicités", "Pour votre information et éventuelles instructions.", False),
    ]

    for i, (label, value, tall) in enumerate(rows_data):
        row = table.rows[i]

        lc = row.cells[0]
        set_col_width(lc, 3.5)
        set_cell_border(lc)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

        vc = row.cells[1]
        set_col_width(vc, 13.0)
        set_cell_border(vc)
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(10)

        if tall:
            trPr = row._tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '2000')
            trHeight.set(qn('w:hRule'), 'atLeast')
            trPr.append(trHeight)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def parse_json_from_text(text: str) -> dict:
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        try:
            raw = json.loads(match.group())

            def flatten(val):
                if isinstance(val, str):
                    return val
                elif isinstance(val, list):
                    return ", ".join(flatten(v) for v in val)
                elif isinstance(val, dict):
                    return " | ".join(f"{flatten(v)}" for v in val.values())
                return str(val)

            return {
                "cadre":            flatten(raw.get("cadre", "")),
                "objet":            flatten(raw.get("objet", "")),
                "participants":     flatten(raw.get("participants", "")),
                "descriptif":       flatten(raw.get("descriptif", "")),
                "prochaine_action": flatten(raw.get("prochaine_action", "")),
            }
        except json.JSONDecodeError:
            pass

    return {
        "cadre": "Non déterminé",
        "objet": "Non déterminé",
        "participants": "Non déterminé",
        "descriptif": text[:300] if text else "Non déterminé",
        "prochaine_action": "Non déterminé"
    }